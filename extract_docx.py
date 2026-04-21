from docx import Document
doc = Document(r'C:\Users\MSIGE79\Downloads\Share administration system requirements.docx')
for p in doc.paragraphs:
    print(p.text)
print("\n\n=== TABLES ===\n")
for i, table in enumerate(doc.tables):
    print(f"\n--- TABLE {i+1} ---")
    for row in table.rows:
        print(" | ".join(cell.text.replace("\n", " ") for cell in row.cells))
